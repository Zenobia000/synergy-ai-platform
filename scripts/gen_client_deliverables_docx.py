"""產生客戶交付清單 docx（給非工程背景人員看）。

來源：docs/13_client_deliverables.md
輸出：docs/to_c_docs/客戶準備清單.docx
規格：窄邊界、12pt、微軟正黑體、白話用語
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "微軟正黑體"
FONT_SIZE = Pt(12)


def set_run_font(run, size=FONT_SIZE, bold=False, color=None):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    sizes = {1: Pt(20), 2: Pt(16), 3: Pt(14)}
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, Pt(14)), bold=True,
                 color=RGBColor(0x1F, 0x4E, 0x79))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    if widths:
        for i, w in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # 標題列
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(hdr[i], "1F4E79")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 資料列
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def main():
    doc = Document()

    # 預設樣式
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = FONT_SIZE
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT)
    rPr.append(rFonts)

    # 窄邊界（上下左右各 1.27 cm，Word「窄」預設值）
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    # === 封面 ===
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Synergy AI 教練成交副駕駛")
    set_run_font(run, size=Pt(24), bold=True,
                 color=RGBColor(0x1F, 0x4E, 0x79))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("客戶準備清單")
    set_run_font(run, size=Pt(18), bold=True)

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_p.add_run("給品牌方 / 業主的準備指南")
    set_run_font(run, size=Pt(14),
                 color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()

    # === 一句話介紹 ===
    add_heading(doc, "這份文件在說什麼？", 2)
    add_para(doc,
        "我們要幫您打造一個「教練成交副駕駛」系統，讓教練在跟客戶溝通時更有效率、"
        "成交率更高、跟進不會漏掉。"
    )
    add_para(doc,
        "這份清單列出「只有您（品牌方）才能提供」的東西，因為這些是您的專業知識"
        "（產品、話術、合規、法規等）。技術部分（雲端、伺服器、資料庫等）我們會"
        "全部處理好，您不用煩惱。"
    )

    # === 6 大準備項目總覽 ===
    add_heading(doc, "您要準備的 6 樣東西", 2)
    add_para(doc, "簡單來說，請您準備好以下六樣，越早越好：", bold=True)

    add_table(doc,
        headers=["#", "您要準備的東西", "什麼時候要交", "為什麼需要"],
        rows=[
            ["1", "客戶狀態判斷邏輯", "第一週週五前", "讓系統知道客戶走到哪一步"],
            ["2", "問卷題目跟計分方式", "第一週週三前 ★急", "讓 AI 看懂客戶的健康狀況"],
            ["3", "商品清單（至少 20 樣）", "第一週週三前 ★急", "讓 AI 知道有什麼可以推薦"],
            ["4", "不能用的話術詞彙", "第一週週三前 ★急", "避免踩到法規地雷"],
            ["5", "標準話術範本", "第一週週五前", "讓 AI 寫出來像您家教練"],
            ["6", "相關法規依據", "第二週週一前", "讓系統合規上線"],
        ],
        widths=[1, 5, 3, 6],
    )

    add_para(doc, "")
    add_para(doc,
        "★ 標記的是「最急的」— 如果這三樣沒到位，第二週開發會卡住。",
        bold=True
    )

    # === 第 1 項 ===
    add_heading(doc, "1. 客戶狀態判斷邏輯", 2)

    add_para(doc, "為什麼需要：", bold=True)
    add_para(doc,
        "系統需要知道一個客戶現在走到哪個階段（剛認識？已經填過問卷？已經談過？還"
        "在猶豫？...），才能對應提醒教練該做什麼。"
    )

    add_para(doc, "請您告訴我們：", bold=True)
    add_bullet(doc, "客戶從第一次接觸到成交，會經過哪些階段？")
    add_bullet(doc, "怎麼判斷客戶從一個階段「走到」下一個階段？")
    add_bullet(doc, "什麼情況算「成交」？（付款？開卡？第一次出貨？）")
    add_bullet(doc, "什麼情況算「失敗」？什麼算「還在等他回」？")
    add_bullet(doc, "有沒有試用期？多久？")

    add_para(doc, "範例（給您參考的格式，您可以直接畫流程圖或寫 Word 檔）：",
             bold=True)
    add_para(doc,
        "「客戶填完問卷 → 教練約商談 → 商談完成 → 推薦商品 → "
        "(試用 / 直接買 / 拒絕)。如果推薦後 7 天沒回應就標『需回訪』，"
        "30 天都沒互動就標『沉默』。」"
    )
    add_para(doc, "格式不限，重點是邏輯講清楚。階段不一定要 10 個，6-10 個都可以。")

    # === 第 2 項 ===
    add_heading(doc, "2. 問卷題目跟計分方式 ★急", 2)

    add_para(doc, "為什麼需要：", bold=True)
    add_para(doc, "這是整個系統的入口。客戶填問卷後，AI 才能判斷他的健康狀況、推什麼商品。")

    add_para(doc, "問卷分三個階段：", bold=True)
    add_table(doc,
        headers=["階段", "題數", "問什麼"],
        rows=[
            ["第一階段：快速分流", "5 題", "年齡、生活型態、最在意什麼"],
            ["第二階段：六大核心", "12 題", "壓力、睡眠、消化、手腳冰冷、水腫、排便"],
            ["第三階段：動態追問", "3-8 題", "依照前面回答，再深入問飲食 / 活動 / 補充品 / 女性週期等"],
        ],
        widths=[5, 2, 8],
    )

    add_para(doc, "")
    add_para(doc, "請您準備：", bold=True)
    add_bullet(doc, "三個階段的所有題目原文（中文）")
    add_bullet(doc, "每題的選項（單選 / 多選 / 填空）")
    add_bullet(doc, "每題對應到哪個健康關注（例：選 A → 壓力相關）")
    add_bullet(doc, "計分方式（每題 0-3 分？怎麼加總？）")
    add_bullet(doc, "什麼答案組合會「亮紅燈」需要提醒教練（例：胸痛 + 家族病史）")
    add_bullet(doc, "第三階段「依回答動態出題」的規則（例：壓力分數高 → 加問壓力相關）")

    add_para(doc, "另外請給我們：", bold=True)
    add_bullet(doc,
        "2-3 篇真實的「客戶版健康摘要」範例（去除個資即可），讓 AI 學您家的口吻"
    )
    add_bullet(doc,
        "2-3 篇「教練版商談筆記」範例（含建議推哪些商品、怎麼切入）"
    )

    # === 第 3 項 ===
    add_heading(doc, "3. 商品清單 ★急", 2)

    add_para(doc, "為什麼需要：", bold=True)
    add_para(doc, "AI 要根據客戶問卷推薦商品，所以需要您家完整的商品資料。")

    add_para(doc, "格式：", bold=True)
    add_para(doc, "Excel 或 Google 試算表都可以。試行階段最少 20 樣商品。")

    add_para(doc, "每樣商品請填：", bold=True)
    add_table(doc,
        headers=["欄位", "必填", "範例"],
        rows=[
            ["商品代碼", "是", "SYN-PROBIO-30"],
            ["中文名稱", "是", "益生菌複方膠囊"],
            ["主要成分", "是", "5 株益生菌、益生元 FOS"],
            ["針對什麼問題", "是", "消化、便秘、免疫"],
            ["建議怎麼吃", "是", "每天 1 顆，餐後溫水"],
            ["禁忌或注意事項", "是", "免疫抑制者請先諮詢醫師"],
            ["不適合搭配", "可選", "與抗生素間隔 2 小時"],
            ["價格（NTD）", "是", "1,200"],
            ["規格", "是", "30 顆 / 瓶"],
            ["賣點摘要（50 字內）", "是", "幫助消化道菌群平衡..."],
            ["不可以宣稱什麼", "可選", "不可說『治療便秘』"],
        ],
        widths=[3.5, 1.5, 7],
    )

    add_para(doc, "")
    add_para(doc,
        "重點：「針對什麼問題」一定要對應到問卷裡的健康標籤，AI 才能正確推薦。",
        bold=True
    )

    # === 第 4 項 ===
    add_heading(doc, "4. 不能用的話術詞彙 ★急", 2)

    add_para(doc, "為什麼需要：", bold=True)
    add_para(doc,
        "AI 寫出來的文案如果亂講「治療糖尿病」「保證月入 10 萬」之類的話，會踩"
        "到法規地雷。系統需要您家的合規團隊（或法務）提供「不能用的詞」清單，"
        "AI 才能自動避開或改寫。"
    )

    add_para(doc, "分四大類，每類至少 40 個詞，總共 200 個以上：", bold=True)
    add_table(doc,
        headers=["類別", "範例", "至少幾個"],
        rows=[
            ["醫療宣稱類", "治療、治癒、預防 + 病名（糖尿病、高血壓、癌症...）", "60 個"],
            ["收入宣稱類", "月入 X 萬、保證收入、被動收入、財富自由", "40 個"],
            ["誇大效果類", "100% 有效、立即見效、神奇療效", "40 個"],
            ["金字塔風險類", "拉人頭、上線抽成、層級獎金", "40 個"],
        ],
        widths=[3, 9, 2],
    )

    add_para(doc, "")
    add_para(doc, "格式（Excel 或 CSV 檔，欄位如下）：", bold=True)
    add_table(doc,
        headers=["類別", "禁用詞", "嚴重度", "可以改成什麼"],
        rows=[
            ["醫療類", "治療糖尿病", "高", "支持血糖管理"],
            ["醫療類", "預防癌症", "高", "強化日常保健"],
            ["收入類", "月入十萬", "高", "額外收入機會"],
            ["誇大類", "百分百有效", "中", "有效成分"],
            ["金字塔類", "拉人頭", "高", "推薦合作夥伴"],
        ],
        widths=[2.5, 4, 2, 5],
    )

    add_para(doc, "")
    add_para(doc, "我們會做的事：", bold=True)
    add_bullet(doc, "把您給的詞庫匯入系統")
    add_bullet(doc, "AI 自動用「智慧比對」抓出改寫變體（例如「根治」也會被抓到）")
    add_bullet(doc,
        "之後可以在後台自己增刪改詞，不用重新部署系統"
    )

    # === 第 5 項 ===
    add_heading(doc, "5. 標準話術範本", 2)

    add_para(doc, "為什麼需要：", bold=True)
    add_para(doc,
        "AI 寫商談話術時，需要參考您家既有的標準話術，才不會「不像 Synergy 的人在說話」。"
    )

    add_para(doc, "請整理以下五類：", bold=True)
    add_table(doc,
        headers=["類別", "內容", "至少幾條"],
        rows=[
            ["開場話術", "不同場景的開場（LINE 邀約 / 活動 / 朋友介紹）", "5 種"],
            ["產品銜接話術", "從問卷痛點接到商品推薦的橋接句", "10 條"],
            ["異議處理話術", "客戶說「太貴」「我再想想」「老公不同意」怎麼回", "15 對"],
            ["柔性成交話術", "不施壓的下單邀約", "5 種"],
            ["跟進話術", "48 小時 / 7 天 / 30 天三個時間點的跟進文案", "每個 3 種"],
        ],
        widths=[3, 7.5, 2],
    )

    add_para(doc, "")
    add_para(doc, "格式建議（讓 AI 更會用）：", bold=True)
    add_para(doc,
        "每條話術標三個標籤：適用情境、對應商品、關鍵字。例如："
    )
    add_para(doc,
        "「客戶說太貴」/「通用」/「價格、CP值」 → 我每天少喝一杯飲料的錢..."
    )

    add_para(doc, "格式不限（Excel / Word / Google Doc 都可），標籤清楚就好。")

    # === 第 6 項 ===
    add_heading(doc, "6. 相關法規依據", 2)

    add_para(doc, "為什麼需要：", bold=True)
    add_para(doc, "讓系統的合規檢查有「法律依據」，未來被質疑時拿得出來說明。")

    add_para(doc, "請您準備：", bold=True)
    add_bullet(doc, "個資法相關聲明（給客戶看的隱私政策草案，我方可協助潤稿）")
    add_bullet(doc, "健康食品管理法 — 哪些用語不可宣稱（衛福部公告對照）")
    add_bullet(doc, "多層次傳銷管理法 — 哪些招募話術不能用（公平會公告）")
    add_bullet(doc, "化妝品管理法（如果有美妝商品）")
    add_bullet(doc, "您家內部如果有比法規更嚴的合規 SOP，也一併提供")

    # === 我方處理事項 ===
    add_heading(doc, "您不用煩惱的事（我方包辦）", 2)
    add_para(doc, "以下技術相關事項全部由我方處理：", bold=True)

    add_table(doc,
        headers=["項目", "我方處理", "需您配合一次的事"],
        rows=[
            ["雲端伺服器（GCP）", "我方代開帳號 + 維運", "若您要自己付月費，需簽署"],
            ["LINE 官方帳號申請", "我方協助申請流程", "提供商家證明文件 + 收驗證信"],
            ["WhatsApp 商用帳號", "我方協助設定", "提供商家證明 + 企業電話"],
            ["寄信網域", "我方用內部網域", "若要用您家網域才需設定 DNS"],
            ["管理員帳號", "我方建立初始帳號", "提供 1 位窗口的 email"],
            ["教練培訓", "我方做 30 分鐘 onboarding", "您方 PM 出席即可"],
            ["客戶資料儲存", "我方加密保存", "簽署資料處理協議（DPA）"],
        ],
        widths=[4, 5, 5],
    )

    add_para(doc, "")
    add_para(doc,
        "唯一還是需要您指派的「人」：試行階段的 3-5 位教練。",
        bold=True
    )

    # === 時程與里程碑 ===
    add_heading(doc, "整體時程簡表", 2)
    add_table(doc,
        headers=["時間", "里程碑", "您要簽核什麼"],
        rows=[
            ["第 0 週 第 1 天", "啟動會", "確認試行教練名單、6 大準備項目排程"],
            ["第 0 週 第 5 天", "需求凍結", "上面 1-5 項全部交付完畢"],
            ["第 3 週 第 5 天", "Beta 驗證", "教練回饋彙整、是否進第 4 週上線"],
            ["第 4 週 第 5 天", "正式上線", "上線同意書、緊急聯絡人"],
        ],
        widths=[3.5, 4, 6.5],
    )

    # === 待確認問題清單 ===
    add_heading(doc, "我們需要您回覆的問題", 2)
    add_table(doc,
        headers=["編號", "問題", "什麼時候要回"],
        rows=[
            ["Q1", "試行的 3-5 位教練人選", "第 0 週 第 1 天"],
            ["Q2", "客戶狀態判斷邏輯（第 1 項）", "第 0 週 第 5 天"],
            ["Q3", "問卷題目 + 計分（第 2 項）★急", "第 0 週 第 3 天"],
            ["Q4", "商品清單 ≥ 20 樣（第 3 項）★急", "第 0 週 第 3 天"],
            ["Q5", "不能用的詞庫 ≥ 200 個（第 4 項）★急", "第 0 週 第 3 天"],
            ["Q6", "標準話術初版（第 5 項）", "第 0 週 第 5 天"],
            ["Q7", "法規依據文件（第 6 項）", "第 1 週 第 1 天"],
            ["Q8", "您方窗口 email（系統管理員用）", "第 0 週 第 1 天"],
        ],
        widths=[2, 9, 3],
    )

    # === 結尾 ===
    add_heading(doc, "有任何問題？", 2)
    add_para(doc, "如果這份清單有看不懂的地方，或您家有特殊狀況需要討論，請隨時聯繫我們。")
    add_para(doc, "")
    add_para(doc, "聯絡窗口", bold=True)
    add_para(doc, "專案經理：kuanwei")
    add_para(doc, "Email：bheadwei0910@gmail.com")
    add_para(doc, "")
    add_para(doc,
        "—— 期待跟您一起，把教練的成交率推上去 ——",
        bold=True
    )

    # === 儲存 ===
    output = "docs/to_c_docs/客戶準備清單.docx"
    doc.save(output)
    print(f"OK: {output}")


if __name__ == "__main__":
    main()
