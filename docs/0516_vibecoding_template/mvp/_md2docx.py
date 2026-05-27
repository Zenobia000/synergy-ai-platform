"""Convert prd.md to prd.docx with narrow margins, 12pt 微軟正黑體."""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
SRC = HERE / "prd.md"
DST = HERE / "prd.docx"

FONT_NAME = "Microsoft JhengHei"  # 微軟正黑體
FONT_SIZE = Pt(12)

doc = Document()

# Narrow margins
for section in doc.sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

# Default style
style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = FONT_SIZE
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)


def _apply_font(run, *, bold=False, size=None, color=None):
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = size or FONT_SIZE
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline_runs(paragraph, text: str, *, base_bold=False, size=None):
    # Strip simple link syntax [text](url) -> text
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _apply_font(run, bold=True, size=size)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            _apply_font(run, bold=base_bold, size=size)
        else:
            run = paragraph.add_run(part)
            _apply_font(run, bold=base_bold, size=size)


def add_heading(text: str, level: int):
    sizes = {1: Pt(20), 2: Pt(16), 3: Pt(14), 4: Pt(13)}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    add_inline_runs(p, text, base_bold=True, size=sizes.get(level, Pt(13)))


def add_paragraph(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text)


def add_bullet(text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(0)
    add_inline_runs(p, text)


def add_quote(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text)


def split_table_row(line: str):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _write_cell_multiline(cell, text: str, *, bold=False, size=Pt(11)):
    """Write text into a cell, splitting on <br> / <br/> / <br /> into separate paragraphs."""
    cell.text = ""
    segments = re.split(r"<br\s*/?>", text, flags=re.IGNORECASE)
    for k, seg in enumerate(segments):
        p = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_inline_runs(p, seg.strip(), base_bold=bold, size=size)


def add_table(rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    rows = [r + [""] * (ncols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            _write_cell_multiline(table.cell(i, j), cell_text, bold=(i == 0), size=Pt(11))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def is_table_separator(line: str) -> bool:
    s = line.strip().strip("|").strip()
    return bool(s) and all(re.fullmatch(r":?-+:?", c.strip()) for c in s.split("|"))


def _set_cell_shading(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_write(cell, text, *, bold=False, size=Pt(11), align=None, shading=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    add_inline_runs(p, text, base_bold=bold, size=size)
    if shading:
        _set_cell_shading(cell, shading)


def render_architecture():
    """Draw a clean 3-layer architecture diagram using shaded tables."""
    center = WD_ALIGN_PARAGRAPH.CENTER

    # Layer 1: Frontend
    t1 = doc.add_table(rows=2, cols=1)
    t1.style = "Table Grid"
    _cell_write(t1.cell(0, 0), "Frontend：Mobile PWA（React 19 + Vite + Tailwind v4）",
                bold=True, size=Pt(12), align=center, shading="DCE6F1")
    _cell_write(t1.cell(1, 0),
                "O1 Today's 5 首頁　|　Contact 詳情頁　|　O5 語音預覽下載　|　"
                "O6 異議卡　|　O8 客戶端問卷　|　四種輔助捕捉 UI（貼上 / OCR / 語音 / Share Sheet）",
                size=Pt(11), align=center, shading="EAF1F8")

    p = doc.add_paragraph()
    p.alignment = center
    p.paragraph_format.space_after = Pt(0)
    add_inline_runs(p, "▼  HTTPS REST + SSE streaming", base_bold=True, size=Pt(11))

    # Layer 2: Backend (12 agents in 3x4 grid)
    agents = [
        ("Contact Insight (C1)", "Sonnet 4.6"),
        ("Life-event Radar (C2)", "規則 + Haiku"),
        ("Mood Sensor (C3 lite)", "Haiku 分類"),
        ("Too-Salesy (C4 lite)", "規則引擎"),
        ("Daily Planner (O1)", "規則 v1"),
        ("Draft Agent (O2)", "Haiku → Sonnet + streaming"),
        ("Sample Follow-up (O3)", "Haiku + 排程"),
        ("TTS Pipeline (O5)", "OpenAI / ElevenLabs"),
        ("Objection Handler (O6)", "Haiku"),
        ("Questionnaire (O8)", "Sonnet 摘要"),
        ("Funnel Agent (G1)", "Sonnet"),
        ("Compliance Whisper (G8)", "regex < 50ms"),
    ]
    rows, cols = 4, 3
    t2 = doc.add_table(rows=rows + 1, cols=cols)
    t2.style = "Table Grid"
    # Header
    hdr = t2.rows[0].cells
    hdr[0].merge(hdr[1]).merge(hdr[2])
    _cell_write(t2.cell(0, 0), "Backend：FastAPI (uv) / Python 3.12　—　12 個 Agent",
                bold=True, size=Pt(12), align=center, shading="E2EFDA")
    for idx, (name, model) in enumerate(agents):
        r = 1 + idx // cols
        c = idx % cols
        cell = t2.cell(r, c)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.alignment = center
        add_inline_runs(p1, name, base_bold=True, size=Pt(11))
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.alignment = center
        add_inline_runs(p2, model, size=Pt(10))
        _set_cell_shading(cell, "F2F8EC")

    p = doc.add_paragraph()
    p.alignment = center
    p.paragraph_format.space_after = Pt(0)
    add_inline_runs(p, "▼", base_bold=True, size=Pt(11))

    # Layer 3: Data / External (3 columns)
    t3 = doc.add_table(rows=2, cols=3)
    t3.style = "Table Grid"
    headers = ["資料層", "AI 模型層", "可觀測 / 合規"]
    bodies = [
        "PostgreSQL + pgvector\nSupabase RLS\n（brand_id + rep_id 多租戶隔離）",
        "Anthropic API\n（Haiku / Sonnet 4.6）\n+ OpenAI TTS\nModel Router\n（≤ $0.30 / rep / day）",
        "ComplianceLog + EventLog\nOTel / Langfuse / Sentry",
    ]
    fills = ["FFF2CC", "FCE4D6", "EDEDED"]
    for j, (h, b, f) in enumerate(zip(headers, bodies, fills)):
        _cell_write(t3.cell(0, j), h, bold=True, size=Pt(12), align=center, shading=f)
        cell = t3.cell(1, j)
        cell.text = ""
        for k, line in enumerate(b.split("\n")):
            p = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.alignment = center
            add_inline_runs(p, line, size=Pt(11))
        # lighter tint for body
        light = {"FFF2CC": "FFF9E6", "FCE4D6": "FDF1EA", "EDEDED": "F7F7F7"}[f]
        _set_cell_shading(cell, light)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def process_lines(lines: list[str]):
    i = 0
    in_arch_section = False
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            j = i + 1
            block_lines = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                block_lines.append(lines[j].rstrip("\n"))
                j += 1
            # In section 5, replace ASCII diagram with structured render
            if in_arch_section:
                render_architecture()
            else:
                # Fallback: render as indented monospace-ish paragraph
                for bl in block_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.left_indent = Cm(0.5)
                    add_inline_runs(p, bl, size=Pt(10))
            i = j + 1
            continue

        # Horizontal rule
        if re.fullmatch(r"-{3,}", stripped):
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            title = m.group(2)
            add_heading(title, level)
            if level == 2:
                in_arch_section = title.lstrip().startswith("5.")
            i += 1
            continue

        # Table: line starts with | and next line is separator
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = split_table_row(line)
            i += 2  # skip separator
            body = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                body.append(split_table_row(lines[i]))
                i += 1
            add_table([header] + body)
            continue

        # Block quote
        if stripped.startswith(">"):
            add_quote(stripped.lstrip(">").strip())
            i += 1
            continue

        # Bullet list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            level = indent // 2
            add_bullet(m.group(2), level=level)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(0)
            add_inline_runs(p, m.group(2))
            i += 1
            continue

        # Plain paragraph (collect continuation lines until blank)
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip("\n")
            if not nxt.strip():
                break
            if re.match(r"^(#{1,6}\s|\||>|\s*[-*]\s|\s*\d+\.\s)", nxt):
                break
            buf.append(nxt.strip())
            i += 1
        add_paragraph(" ".join(buf))


text = SRC.read_text(encoding="utf-8")
process_lines(text.splitlines())
doc.save(DST)
print(f"Saved: {DST}")
