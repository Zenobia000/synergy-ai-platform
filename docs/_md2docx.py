"""Generic Markdown -> Word(.docx) converter.

用法：
    python _md2docx.py <markdown檔路徑> [輸出.docx路徑]
不指定輸出時，輸出與來源同名的 .docx。

格式：窄邊界（1.27cm）、內文/表格 12pt、微軟正黑體（Microsoft JhengHei）。
支援：標題、段落、項目符號（含巢狀）、編號清單、表格（含 <br> 換行）、
      引言、水平線、行內粗體 **、行內 code `，並把 [文字](連結) 還原成純文字。
"""
from __future__ import annotations
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if len(sys.argv) < 2:
    raise SystemExit("用法：python _md2docx.py <markdown檔> [輸出.docx]")

SRC = Path(sys.argv[1])
DST = Path(sys.argv[2]) if len(sys.argv) > 2 else SRC.with_suffix(".docx")

FONT_NAME = "Microsoft JhengHei"  # 微軟正黑體
BODY_SIZE = Pt(12)
TABLE_SIZE = Pt(12)

doc = Document()

# 窄邊界（Word「窄」= 1.27 cm 四邊）
for section in doc.sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

# 預設樣式：微軟正黑體 12pt
style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = BODY_SIZE
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)


def _apply_font(run, *, bold=False, size=None):
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = size or BODY_SIZE
    if bold:
        run.bold = True


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline_runs(paragraph, text: str, *, base_bold=False, size=None):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # [文字](連結) -> 文字
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _apply_font(paragraph.add_run(part[2:-2]), bold=True, size=size)
        elif part.startswith("`") and part.endswith("`"):
            _apply_font(paragraph.add_run(part[1:-1]), bold=base_bold, size=size)
        else:
            _apply_font(paragraph.add_run(part), bold=base_bold, size=size)


def add_heading(text: str, level: int):
    sizes = {1: Pt(20), 2: Pt(16), 3: Pt(14), 4: Pt(13)}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    add_inline_runs(p, text, base_bold=True, size=sizes.get(level, Pt(13)))


def add_paragraph(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text)


def add_bullet(text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(0)
    add_inline_runs(p, text)


def add_quote(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text)


def split_table_row(line: str):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _write_cell_multiline(cell, text: str, *, bold=False, size=TABLE_SIZE):
    cell.text = ""
    for k, seg in enumerate(re.split(r"<br\s*/?>", text, flags=re.IGNORECASE)):
        p = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_inline_runs(p, seg.strip(), base_bold=bold, size=size)


def add_table(rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    rows = [r + [""] * (ncols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            _write_cell_multiline(table.cell(i, j), cell_text, bold=(i == 0), size=TABLE_SIZE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def is_table_separator(line: str) -> bool:
    s = line.strip().strip("|").strip()
    return bool(s) and all(re.fullmatch(r":?-+:?", c.strip()) for c in s.split("|"))


def process_lines(lines: list[str]):
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 圍籬程式碼區塊 -> 縮排呈現
        if stripped.startswith("```"):
            j = i + 1
            block_lines = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                block_lines.append(lines[j].rstrip("\n"))
                j += 1
            for bl in block_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Cm(0.5)
                add_inline_runs(p, bl, size=Pt(11))
            i = j + 1
            continue

        # 水平線
        if re.fullmatch(r"[-*_]{3,}", stripped):
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # 標題
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            add_heading(m.group(2), len(m.group(1)))
            i += 1
            continue

        # 表格（本行以 | 開頭，下一行為分隔線）
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = split_table_row(line)
            i += 2
            body = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                body.append(split_table_row(lines[i]))
                i += 1
            add_table([header] + body)
            continue

        # 引言
        if stripped.startswith(">"):
            add_quote(stripped.lstrip(">").strip())
            i += 1
            continue

        # 項目符號（含巢狀）
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            add_bullet(m.group(2), level=len(m.group(1)) // 2)
            i += 1
            continue

        # 編號清單
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(0)
            add_inline_runs(p, m.group(2))
            i += 1
            continue

        # 一般段落（收集連續行直到空行/區塊起始）
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip("\n")
            if not nxt.strip():
                break
            if re.match(r"^(#{1,6}\s|\||>|\s*[-*]\s|\s*\d+\.\s|```)", nxt):
                break
            buf.append(nxt.strip())
            i += 1
        add_paragraph(" ".join(buf))


process_lines(SRC.read_text(encoding="utf-8").splitlines())
doc.save(DST)
print(f"Saved: {DST}")
